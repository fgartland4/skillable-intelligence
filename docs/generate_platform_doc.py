#!/usr/bin/env python3
"""
Generate Skillable Intelligence Platform briefing Word document.
Audience: Skillable executive leadership team.
Structure: Three Problems Worth Solving / Three Tools. One Platform. / HubSpot / Recommendations
"""

import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

DOCS_DIR    = os.path.dirname(os.path.abspath(__file__))
OUTPUT_PATH = os.path.join(DOCS_DIR, "Skillable-Intelligence-Platform.docx")
LOGO_PATH   = r"C:\Users\Frank.Gartland\OneDrive - Skillable\Sales Enablement\Keep\z-Skillable Logos\Skillable Logo\Default@4x.png"
ICON_PATH   = os.path.join(DOCS_DIR, "ai-moment-icon.png")

# ── Colors — Skillable brand palette (confirmed 2026-03-31) ──────────────────
DARK_GREEN     = RGBColor(0x13, 0x69, 0x45)   # #136945
DARK_GREEN_HEX = "136945"
PURPLE         = RGBColor(0x70, 0x00, 0xFF)   # #7000FF
PURPLE_HEX     = "7000FF"
DARK_TEXT      = RGBColor(0x1A, 0x1A, 0x1A)
WHITE          = RGBColor(0xFF, 0xFF, 0xFF)
GRAY           = RGBColor(0x60, 0x60, 0x60)
ROW_ALT        = "F0F5F2"
FONT_NAME      = "Calibri"

# ── Icon ID counter (unique per document) ─────────────────────────────────────
_icon_id = [100]


# ── Typographic quotes ────────────────────────────────────────────────────────

def smartify(text):
    """Convert straight quotes to typographic (curly) quotes."""
    import re
    text = re.sub(r'(?<=[(\s\u2014])"(?=\S)', '\u201c', text)
    text = re.sub(r'^"(?=\S)', '\u201c', text)
    text = re.sub(r'"(?=[\s.,;:!?)\u2014]|$)', '\u201d', text)
    text = re.sub(r'"', '\u201c', text)
    text = re.sub(r"(?<=\w)'(?=\w)", '\u2019', text)
    text = re.sub(r"(?<=\s)'(?=\S)", '\u2018', text)
    text = re.sub(r"'(?=[\s.,;:!?]|$)", '\u2019', text)
    return text


# ── XML / formatting helpers ──────────────────────────────────────────────────

def set_cell_background(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def set_cell_margins(cell, val=50):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement("w:tcMar")
    for side in ["top", "left", "bottom", "right"]:
        m = OxmlElement(f"w:{side}")
        m.set(qn("w:w"), str(val))
        m.set(qn("w:type"), "dxa")
        tcMar.append(m)
    tcPr.append(tcMar)


def set_cell_width(cell, dxa):
    tcPr = cell._tc.get_or_add_tcPr()
    tcW = OxmlElement("w:tcW")
    tcW.set(qn("w:w"), str(dxa))
    tcW.set(qn("w:type"), "dxa")
    tcPr.append(tcW)


def set_paragraph_spacing(paragraph, before=0, after=0, line=None):
    pPr = paragraph._p.get_or_add_pPr()
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:before"), str(before))
    spacing.set(qn("w:after"), str(after))
    if line is not None:
        spacing.set(qn("w:line"), str(line))
        spacing.set(qn("w:lineRule"), "auto")
    pPr.append(spacing)


def add_bottom_border(paragraph, color_hex=DARK_GREEN_HEX, size=8):
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color_hex)
    pBdr.append(bottom)
    pPr.append(pBdr)


def make_run(paragraph, text, size_pt=10, bold=False, italic=False, color=None):
    run = paragraph.add_run(smartify(text))
    run.font.name = FONT_NAME
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color
    return run


def page_break(doc):
    p = doc.add_paragraph()
    set_paragraph_spacing(p, 0, 0)
    run = p.add_run()
    br = OxmlElement("w:br")
    br.set(qn("w:type"), "page")
    run._r.append(br)


# ── Document setup ────────────────────────────────────────────────────────────

def setup_document():
    doc = Document()
    section = doc.sections[0]
    section.page_width    = Inches(8.5)
    section.page_height   = Inches(11)
    section.left_margin   = Inches(0.85)
    section.right_margin  = Inches(0.85)
    section.top_margin    = Inches(0.9)
    section.bottom_margin = Inches(0.9)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)
    return doc, section


# ── Header ────────────────────────────────────────────────────────────────────

def add_header(section):
    header = section.header
    header.is_linked_to_previous = False
    for p in header.paragraphs:
        p._element.getparent().remove(p._element)
    p = header.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_paragraph_spacing(p, 0, 0)
    run = p.add_run()
    run.add_picture(LOGO_PATH, width=Inches(1.1))


# ── Footer ────────────────────────────────────────────────────────────────────

def page_num_fld(paragraph, instrText):
    fldChar_begin = OxmlElement("w:fldChar")
    fldChar_begin.set(qn("w:fldCharType"), "begin")
    r1 = OxmlElement("w:r")
    rPr1 = OxmlElement("w:rPr")
    rFonts1 = OxmlElement("w:rFonts"); rFonts1.set(qn("w:ascii"), FONT_NAME); rFonts1.set(qn("w:hAnsi"), FONT_NAME)
    sz1 = OxmlElement("w:sz"); sz1.set(qn("w:val"), "16")
    color1 = OxmlElement("w:color"); color1.set(qn("w:val"), "888888")
    rPr1.extend([rFonts1, sz1, color1])
    r1.extend([rPr1, fldChar_begin])

    r2 = OxmlElement("w:r")
    rPr2 = OxmlElement("w:rPr")
    rFonts2 = OxmlElement("w:rFonts"); rFonts2.set(qn("w:ascii"), FONT_NAME); rFonts2.set(qn("w:hAnsi"), FONT_NAME)
    sz2 = OxmlElement("w:sz"); sz2.set(qn("w:val"), "16")
    color2 = OxmlElement("w:color"); color2.set(qn("w:val"), "888888")
    rPr2.extend([rFonts2, sz2, color2])
    instr = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve"); instr.text = instrText
    r2.extend([rPr2, instr])

    fldChar_end = OxmlElement("w:fldChar")
    fldChar_end.set(qn("w:fldCharType"), "end")
    r3 = OxmlElement("w:r")
    rPr3 = OxmlElement("w:rPr")
    rFonts3 = OxmlElement("w:rFonts"); rFonts3.set(qn("w:ascii"), FONT_NAME); rFonts3.set(qn("w:hAnsi"), FONT_NAME)
    sz3 = OxmlElement("w:sz"); sz3.set(qn("w:val"), "16")
    color3 = OxmlElement("w:color"); color3.set(qn("w:val"), "888888")
    rPr3.extend([rFonts3, sz3, color3])
    r3.extend([rPr3, fldChar_end])

    paragraph._p.extend([r1, r2, r3])


def gray_run(paragraph, text):
    r = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    rFonts = OxmlElement("w:rFonts"); rFonts.set(qn("w:ascii"), FONT_NAME); rFonts.set(qn("w:hAnsi"), FONT_NAME)
    sz = OxmlElement("w:sz"); sz.set(qn("w:val"), "16")
    color = OxmlElement("w:color"); color.set(qn("w:val"), "888888")
    rPr.extend([rFonts, sz, color])
    t = OxmlElement("w:t"); t.set(qn("xml:space"), "preserve"); t.text = text
    r.extend([rPr, t])
    paragraph._p.append(r)


def add_footer(section):
    footer = section.footer
    footer.is_linked_to_previous = False
    for p in footer.paragraphs:
        p._element.getparent().remove(p._element)
    p = footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_paragraph_spacing(p, 0, 0)
    gray_run(p, "Page ")
    page_num_fld(p, " PAGE ")
    gray_run(p, " of ")
    page_num_fld(p, " NUMPAGES ")


# ── Content helpers ───────────────────────────────────────────────────────────

def add_title(doc, title, subtitle):
    p = doc.add_paragraph()
    set_paragraph_spacing(p, 0, 40)
    make_run(p, title, size_pt=20, bold=True, color=DARK_GREEN)
    p2 = doc.add_paragraph()
    set_paragraph_spacing(p2, 0, 100)
    add_bottom_border(p2)
    make_run(p2, subtitle, size_pt=10, italic=False, color=GRAY)


def h1(doc, text):
    p = doc.add_paragraph()
    set_paragraph_spacing(p, 180, 50)
    pPr = p._p.get_or_add_pPr()
    kwn = OxmlElement("w:keepNext"); kwn.set(qn("w:val"), "1"); pPr.append(kwn)
    add_bottom_border(p, size=4)
    make_run(p, text, size_pt=13, bold=True, color=DARK_GREEN)
    return p


def h2(doc, text):
    p = doc.add_paragraph()
    set_paragraph_spacing(p, 140, 40)
    pPr = p._p.get_or_add_pPr()
    kwn = OxmlElement("w:keepNext"); kwn.set(qn("w:val"), "1"); pPr.append(kwn)
    make_run(p, text, size_pt=11, bold=True, color=DARK_GREEN)
    return p


def body(doc, text, after=72):
    p = doc.add_paragraph()
    set_paragraph_spacing(p, 0, after, line=264)
    make_run(p, text, color=DARK_TEXT)
    return p


def body_bold(doc, text_parts, after=72):
    """text_parts: list of (text, bold) tuples — all dark text."""
    p = doc.add_paragraph()
    set_paragraph_spacing(p, 0, after, line=264)
    for text, bold in text_parts:
        make_run(p, text, bold=bold, color=DARK_TEXT)
    return p


def _bullet_para(doc, after=48):
    """Base bullet paragraph — small elegant bullet, tab-aligned text, correct hang."""
    p = doc.add_paragraph()
    set_paragraph_spacing(p, 0, after)
    pPr = p._p.get_or_add_pPr()

    # Remove any inherited indent so ours wins
    existing = pPr.find(qn("w:ind"))
    if existing is not None:
        pPr.remove(existing)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "480")    # wrapped lines align here
    ind.set(qn("w:hanging"), "240") # bullet hangs 240 twips left of text
    pPr.append(ind)

    # Tab stop at the left indent so bullet + tab always aligns text to 480
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "left")
    tab.set(qn("w:pos"), "480")
    tabs.append(tab)
    pPr.append(tabs)

    r = p.add_run("\u2022\t")       # bullet + tab (tab jumps to 480)
    r.font.name = FONT_NAME
    r.font.size = Pt(7)
    r.font.color.rgb = DARK_TEXT
    return p


def bullet(doc, text, after=40):
    p = _bullet_para(doc, after)
    make_run(p, text, color=DARK_TEXT)
    return p


def bullet_bold(doc, label, rest, after=40):
    p = _bullet_para(doc, after)
    make_run(p, label, bold=True, color=DARK_TEXT)
    make_run(p, rest, color=DARK_TEXT)
    return p


def section_note(doc, text, after=60):
    p = doc.add_paragraph()
    set_paragraph_spacing(p, 40, after)
    make_run(p, text, italic=True, color=GRAY)
    return p


def proclamation(doc, text):
    """Bold declarative line — the pivot between problem and solution in each subsection."""
    p = doc.add_paragraph()
    set_paragraph_spacing(p, 90, 90)
    make_run(p, text, size_pt=11, bold=True, color=DARK_GREEN)
    return p


def add_table(doc, headers, rows, col1_dxa=2800):
    """Full-width table: dark green header, alternating rows.
    col1_dxa: width of column 1; column 2 gets the remainder."""
    total_dxa = 9792
    col2_dxa  = total_dxa - col1_dxa
    col_widths = [col1_dxa, col2_dxa]
    col_count  = len(headers)

    table = doc.add_table(rows=1 + len(rows), cols=col_count)
    table.style = "Table Grid"

    # Header row
    hdr_row = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr_row.cells[i]
        set_cell_background(cell, DARK_GREEN_HEX)
        set_cell_margins(cell, 50)
        set_cell_width(cell, col_widths[i] if i < 2 else total_dxa // col_count)
        for p in cell.paragraphs:
            p._element.getparent().remove(p._element)
        p = cell.add_paragraph()
        set_paragraph_spacing(p, 0, 0)
        make_run(p, h, bold=True, color=WHITE)

    # Data rows
    for r_idx, row_data in enumerate(rows):
        row = table.rows[r_idx + 1]
        bg = ROW_ALT if r_idx % 2 == 1 else "FFFFFF"
        for c_idx, cell_text in enumerate(row_data):
            cell = row.cells[c_idx]
            set_cell_background(cell, bg)
            set_cell_margins(cell, 50)
            set_cell_width(cell, col_widths[c_idx] if c_idx < 2 else total_dxa // col_count)
            for p in cell.paragraphs:
                p._element.getparent().remove(p._element)
            p = cell.add_paragraph()
            set_paragraph_spacing(p, 0, 0)
            make_run(p, smartify(cell_text), color=DARK_TEXT)

    p_after = doc.add_paragraph()
    set_paragraph_spacing(p_after, 0, 60)


# ── AI Moment: inline purple words + floating margin icon ─────────────────────

def add_margin_icon(doc, para):
    """Add AI moment chip icon as floating image anchored in the right page margin.
    Uses add_picture() for image registration, then restructures inline -> anchor."""
    from copy import deepcopy

    WP = "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
    A  = "http://schemas.openxmlformats.org/drawingml/2006/main"

    # Let python-docx add the picture inline — handles image part + rId correctly
    tmp_run = para.add_run()
    tmp_run.add_picture(ICON_PATH, width=Inches(0.23))

    drawing = tmp_run._r.find(qn("w:drawing"))
    if drawing is None:
        return

    inline = drawing.find(f"{{{WP}}}inline")
    if inline is None:
        return  # leave as inline if structure unexpected

    # Grab the pieces we need from the inline element
    graphic = inline.find(f"{{{A}}}graphic")
    docPr   = inline.find(f"{{{WP}}}docPr")
    cnvFr   = inline.find(f"{{{WP}}}cNvGraphicFramePr")

    # Remove inline; we'll replace with anchor inside same drawing element
    drawing.remove(inline)

    cx    = cy    = int(0.23 * 914400)           # 0.23" in EMU
    pos_x = int((7.65 + 0.08) * 914400)          # just into right margin from page left
    pos_y = int(0.01 * 914400)                    # tiny vertical offset

    anchor = OxmlElement("wp:anchor")
    for k, v in [("distT","0"),("distB","0"),("distL","114300"),("distR","114300"),
                 ("simplePos","0"),("relativeHeight","251658240"),("behindDoc","0"),
                 ("locked","0"),("layoutInCell","1"),("allowOverlap","1")]:
        anchor.set(k, v)

    sp = OxmlElement("wp:simplePos"); sp.set("x","0"); sp.set("y","0")
    anchor.append(sp)

    ph = OxmlElement("wp:positionH"); ph.set("relativeFrom","page")
    poh = OxmlElement("wp:posOffset"); poh.text = str(pos_x); ph.append(poh)
    anchor.append(ph)

    pv = OxmlElement("wp:positionV"); pv.set("relativeFrom","paragraph")
    pov = OxmlElement("wp:posOffset"); pov.text = str(pos_y); pv.append(pov)
    anchor.append(pv)

    ext = OxmlElement("wp:extent"); ext.set("cx", str(cx)); ext.set("cy", str(cy))
    anchor.append(ext)

    ee = OxmlElement("wp:effectExtent")
    for k in ["l","t","r","b"]: ee.set(k, "0")
    anchor.append(ee)

    anchor.append(OxmlElement("wp:wrapNone"))

    if docPr  is not None: anchor.append(deepcopy(docPr))
    if cnvFr  is not None: anchor.append(deepcopy(cnvFr))
    if graphic is not None: anchor.append(deepcopy(graphic))

    drawing.append(anchor)


def ai_para(doc, text_parts, after=72):
    """Body paragraph with AI moment phrase in bold purple + floating margin icon.
    text_parts: list of (text, is_purple) tuples — purple parts are also bold."""
    p = doc.add_paragraph()
    set_paragraph_spacing(p, 0, after, line=264)
    for text, is_purple in text_parts:
        run = p.add_run(smartify(text))
        run.font.name = FONT_NAME
        run.font.size = Pt(10)
        run.font.color.rgb = PURPLE if is_purple else DARK_TEXT
        run.font.bold = is_purple
    add_margin_icon(doc, p)
    return p


# ── Section: Three Problems Worth Solving ────────────────────────────────────

def write_why(doc):
    h1(doc, "Three problems worth solving")

    body(doc, "We recently agreed that Skillable is PaaS and not SaaS. Does this distinction simply change some positioning or does it signal our opportunity to shatter a few glass ceilings? As the category creator for hands-on experience platforms, we find ourselves staring at three unique challenges.")

    bullet(doc, "Identifying companies in our ICP")
    bullet(doc, "Proving labability and lab impact")
    bullet(doc, "Enabling bought-in customers to adopt labs", after=60)

    h2(doc, "The PaaS Difference")
    body(doc, "Most B2B software companies sell products that run in a cloud, in a datacenter, on a computer, or on a phone. With a few customizations here and there, many of these products are roughly the same for every user. Qualifying a prospect means finding buyers who fit the profile: the right size, the right industry, the right pain point, the right budget.")
    body(doc, "Skillable is different. We do not deliver our own software. We orchestrate other companies\u2019 software into hands-on lab environments \u2014 real products, real interfaces, real consequence of error, accessible to a learner anywhere. What we sell is the platform that makes so many different things possible at scale.")
    body(doc, "This creates a fundamentally different go-to-market challenge. The buyer profile question \u2014 does this company value training? \u2014 is necessary but not sufficient. The deeper question is whether their products can be orchestrated into a Skillable environment at all. Whether those products are complex enough that hands-on practice creates real value. Whether the company can build and sustain a lab program once they are a customer.")
    body(doc, "These questions cannot be answered with the exhaustive-but-typical data provided by leading Marketing account intelligence software. They require a different kind of analysis entirely. That is what Intelligence provides.")


# ── Section: Three Challenges. Three Tools. One Platform. ────────────────────

def write_what(doc):
    h1(doc, "Three challenges. Three tools. One platform.")

    body(doc, "Each tool addresses one of the three challenges directly. All three share the same research and scoring engine \u2014 which means every analysis makes the entire platform smarter. A company Prospector evaluates is available to Inspector without re-running research. An Inspector analysis seeds Designer with the product context it needs from day one. The intelligence compounds with every use.")

    h2(doc, "Designer: Closing the adoption gap")
    body(doc, "A signed contract is not an adopted customer. The gap between \u201cwe want labs\u201d and \u201cwe\u2019re building and delivering labs embedded in a variety of learning journeys\u201d is wide. And it\u2019s crystal clear that most customers are not, and cannot, make the leap without significant and structured support.")
    body(doc, "Program design is a specialized skill. Embedding a new modality like hands-on experiences into your overall content strategy and content development operation is not a decision. It\u2019s a new discipline. They\u2019ve built documentation, recorded videos, written certification exams. They have not mapped a product\u2019s administrative workflows to a sequence of learner activities, defined scoring logic for hands-on tasks, or produced a structured brief that a lab developer can build against without extensive back-and-forth.")
    body(doc, "Without a process that takes them from goals and audience to a complete, buildable program architecture, customers stall at the design phase. In a consumption model, a stalled customer is not just a missed upsell opportunity; it is a churn risk. A customer who has not built a program has not realized value. A customer who has not realized value does not renew.")
    proclamation(doc, "Designer closes that gap.")
    body(doc, "Designer guides program owners, instructional designers, and subject matter experts through the full process \u2014 from learning objectives and intended audience through every decision a program requires. It doesn\u2019t require the customer to know how to design a lab program. Designer asks the right questions and sequences the decisions correctly. Every program produces:")
    bullet(doc, "A structured program outline")
    bullet(doc, "Draft lab instructions for every lab in the program")
    bullet(doc, "Learner activities for progress tracking")
    bullet(doc, "Scoring methodology recommendations", after=60)
    ai_para(doc, [
        ("Designer generates a complete Bill of Materials", True),
        (" from everything it knows about the program and the product: environment templates, PowerShell and Bash scripts, Bicep and CloudFormation templates, lifecycle action scripts, credential pool configuration, scoring validation stubs. What previously required hours of Solutions Engineering and lab developer work is produced in the same session where the program was designed.", False),
    ])
    body_bold(doc, [("What it unlocks: ", True), ("Every new customer engagement starts with Designer in the first week \u2014 while the technical folks are working to get the technical details sorted. Program design and environment build run in parallel. Day one is productive for the program owner. The adoption gap closes before it has a chance to form.", False)])

    h2(doc, "Inspector: Proving labability and impact")
    body(doc, "Proving that a Skillable lab program will work for a specific customer\u2019s products requires deep technical analysis. Which is the best delivery path for each software environment? A set of virtual machines or containers, leveraging Azure or AWS subscriptions, a custom API orchestration, or a hybrid setup? What are the architectural constraints? What would a realistic program look like in terms of scope, seat time, and scoring approach? What is the estimated consumption potential?")
    body(doc, "This is the type of work that\u2019s required for virtually every new lab program with every new and existing customer and demands substantial commitment from our Solution Engineers and TSMs. It takes hours of conversation, researching API documents, and plain old trial and error.")
    body(doc, "The result is that qualification depth is rationed. It flows to deals already far enough along to justify the time. Early-stage prospects get a general conversation. The technical questions that would surface a Workday pattern early (before marketing dollars are spent, before SE time is committed) often go unasked until it is too late.")
    proclamation(doc, "Inspector proves labability and impact.")
    body(doc, "Inspector performs a deep product-level analysis of a specific company.")
    bullet_bold(doc, "The Case Board. ", "A broad scan that surfaces all of a company\u2019s products, ranked by labability, with competitive pairings, company-level signals, and an overall fit score. You walk in the room and get the picture at a glance.")
    bullet_bold(doc, "The Dossier. ", "The seller or SE selects three to four products from the Case Board for exhaustive analysis \u2014 full technical orchestrability evidence, delivery path recommendation with rationale, scoring approach, consumption potential estimate, and program scope.", after=60)
    ai_para(doc, [
        ("Inspector turns every sales conversation into a solution conversation.", True),
        (" A seller walking into a meeting with a Dossier knows what the customer\u2019s products can and cannot do on the Skillable platform, which delivery path makes sense and why, and what the estimated consumption potential is. That is not discovery. That is a standing start.", False),
    ])
    body(doc, "Inspector also surfaces the competitive map for every analyzed company, which feeds directly into Prospector\u2019s lookalike targeting.")
    body_bold(doc, [("What it unlocks: ", True), ("Pre-call preparation that was previously impossible at scale is now standard. Every seller and SE enters every conversation with the technical depth that used to require hours of individual research, applied automatically to every product before the first meeting.", False)])

    h2(doc, "Prospector: Finding the right companies")
    body(doc, "Platform companies cannot qualify prospects the way product companies do. The tools Marketing uses (ZoomInfo, 6sense, HubSpot, LinkedIn Sales Navigator) are built to identify buyers who match a profile. For Skillable, that is the wrong question. The right question is whether a company\u2019s products can be delivered as hands-on lab experiences. That is a technical assessment, not a firmographic one.")
    body(doc, "We evaluate every prospect across three dimensions:")
    bullet_bold(doc, "Can we deliver a lab for this company\u2019s products? ", "This is the primary filter. If the answer is no, nothing else matters \u2014 not the size of their training organization, not the depth of their content team, not their enthusiasm for hands-on learning. A company whose products cannot be orchestrated into a Skillable environment is not a prospect.")
    bullet_bold(doc, "Is the product complex enough for labs to create real value? ", "Simple products with shallow workflows don\u2019t benefit enough from hands-on practice to justify the investment. Products with deep administrative workflows, meaningful configuration decisions, and real consequence of error \u2014 those are where labs change what learners can actually do.")
    bullet_bold(doc, "Does the organization have what it takes to build and sustain a program? ", "Content team skills, technical enablement maturity, program leadership. Some companies have it today. Others have the organizational DNA to build it. Either can become a strong customer. Companies with neither are high-risk regardless of product fit.", after=60)
    body(doc, "The Workday pattern illustrates what happens when this analysis doesn\u2019t happen early. On every traditional marketing signal, Workday is an ideal prospect: world-class training organization, dedicated learning division, deep technical enablement culture, massive install base. Two of the three dimensions are strong. The third ends the conversation.")
    bullet_bold(doc, "Pure multi-tenant architecture: ", "Every customer shares the same cloud environment. There is no Workday instance to give a learner.")
    bullet_bold(doc, "No provisioning API: ", "No mechanism to spin up an individual environment programmatically. Skillable\u2019s entire delivery model depends on this capability.")
    bullet_bold(doc, "No deployment model: ", "Nothing to install, containerize, or slice.", after=60)
    body(doc, "These are specific technical facts findable in public documentation before a single sales conversation begins. Workday wasn\u2019t a bad lead. It was motivated, capable people who invested significant time before hitting a wall that was always there \u2014 because product-level technical fit was never evaluated before the pursuit began.")
    body(doc, "The same logic runs in the other direction. When Fortinet is a strong fit, it\u2019s not because Fortinet resembles other good customers as a company. It\u2019s because Fortinet\u2019s products have specific technical characteristics \u2014 multi-VM topology, deep administrative workflows, strong API surface, real consequence of misconfiguration \u2014 that make them ideal for hands-on labs. Every company selling products with those same characteristics is a strong fit for the same reasons. The competitive map of a strong-fit customer is a pre-qualified prospect list.")
    proclamation(doc, "Prospector finds the right companies.")
    body(doc, "Prospector is the go-to-market tool for Marketing and RevOps. It takes a list of companies and returns a ranked assessment of ICP fit \u2014 with product-level evidence, composite scores, verdicts, delivery path signals, and key contacts for every company on the list.")
    ai_para(doc, [
        ("Prospector qualifies every company on product-level fit", True),
        (" before a sequence is written, a dollar is spent, or an SDR makes a call. Companies that clear all three dimensions go into outreach. Workday patterns come off the list \u2014 with specific, documented technical reasons on the Company record. Prospector also surfaces customer expansion opportunities, mapping the department landscape of existing accounts to identify adoption opportunities for existing labs, greenfield departments, and buyers ready to expand.", False),
    ])
    body_bold(doc, [("What it unlocks: ", True), ("The list that goes into outreach is the right list. Every company on it has been evaluated on the question that actually determines Skillable fit \u2014 not on firmographic proxies that have nothing to do with whether their products can be orchestrated into a lab.", False)])
    body(doc, "Fit scores, product signals, delivery path recommendations, key contacts \u2014 all of it needs to reach the right people at the right moment. That is what the HubSpot integration is built to do.")


# ── Section: HubSpot and Revenue Operations ───────────────────────────────────

def write_how(doc):
    h1(doc, "Surfacing the data. Right place. Right time. Right context.")
    body(doc, "The balance of this document outlines the decisions for RevOps to make with Marketing and Security. It includes more context, several recommendations, and a list of decisions to be made. If you\u2019re not in one of those groups, feel free to stop reading if you so choose.")

    h2(doc, "The Integration Principle")
    body(doc, "HubSpot is the seller and marketer\u2019s workspace. Intelligence is the specialist workspace. The integration surfaces the right intelligence, to the right people, in the right places, with the right context \u2014 without requiring sellers to live in another tool or RevOps to build a parallel system.")

    h2(doc, "Prospector \u2194 HubSpot: Bidirectional, marketing-driven")
    body(doc, "Marketing triggers Prospector from inside HubSpot \u2014 selecting a ZoomInfo list or defining criteria and sending them to Prospector for analysis. Prospector writes enriched data back to HubSpot Company records, Contact records, and Deals. HubSpot is Prospector\u2019s primary output destination.")
    body(doc, "Data written to the Company record:")
    add_table(doc,
        ["Data", "Purpose"],
        [
            ["Intelligence Fit Score",        "Numeric score (0\u2013100); enables list segmentation and prioritization"],
            ["Intelligence Verdict",          "Labable / Simulation Candidate / Do Not Pursue; enables filtered views and enrollment triggers"],
            ["Fit Rationale Summary",         "2\u20133 sentences: why this company scored well or poorly, which product drove the score, what the path looks like"],
            ["Top Product Signal",            "The product or product category that most drove the score"],
            ["Recommended Delivery Path",     "Cloud Slice / Standard VM / Simulation / Custom API"],
            ["Key Risk Flag",                 "The single most important constraint a seller needs to know"],
            ["Date of Last Analysis",         "Enables freshness filtering"],
            ["Link to Full Intelligence Report", "One-click access to complete scoring and evidence"],
        ],
        col1_dxa=2600
    )
    body(doc, "Intelligence surfaces up to two contacts per company \u2014 a decision maker and a day-to-day champion \u2014 extracted from public sources for ABM targeting.")

    h2(doc, "Inspector \u2192 HubSpot: Seller and SE-driven")
    body(doc, "Company and Deal records in HubSpot surface a \u2018Run Inspector\u2019 link. Clicking it opens Inspector at the Case Board. The full Inspector experience runs in Intelligence \u2014 HubSpot is only the trigger.")
    body(doc, "Case Board data written to the Company record:")
    add_table(doc,
        ["Data", "Purpose"],
        [
            ["Product list with labability scores",       "What this company sells and how labable each product is"],
            ["Top delivery path signal per product",      "Quick read on delivery approach per product"],
            ["Overall company fit score",                 "Single number for segmentation and prioritization"],
            ["Key risk flag",                             "The most important constraint a seller needs to know"],
            ["Date of last analysis + link to full report", "Freshness tracking and one-click access"],
        ],
        col1_dxa=4800
    )
    body(doc, "Dossier data written to the Deal record:")
    add_table(doc,
        ["Data", "Purpose"],
        [
            ["Delivery path recommendation + rationale",  "What path, and why"],
            ["Scoring approach recommendation",           "How learner actions will be validated"],
            ["Consumption potential / ACV estimate",      "Deal-level revenue context"],
            ["Technical orchestrability evidence",        "Compressed analysis for the SE"],
            ["Program scope estimate",                    "Labs, seat time, curriculum depth"],
            ["Link to full Inspector Dossier",            "Full context one click away"],
        ],
        col1_dxa=4800
    )

    h2(doc, "Designer \u2192 HubSpot: Read-only visibility")
    body(doc, "Program owners and IDs go directly to Designer. HubSpot plays no role in triggering Designer\u2019s workflow.")
    body(doc, "Designer-created Lab Programs surface in HubSpot as read-only links and summary data on the Company record \u2014 giving sellers and CSMs visibility into what programs have been designed, what is in progress, and what has been delivered. Critical context for renewal and expansion conversations before a QBR.")


# ── Section: Recommendations and Open Decisions ───────────────────────────────

def write_recommendations(doc):
    h1(doc, "Recommendations & open decisions")

    h2(doc, "What we are recommending")
    bullet_bold(doc, "Prospector as the primary Marketing data source for ICP outbound. ", "Replace or supplement the current ZoomInfo-only scoring motion with Intelligence-qualified lists. Companies that fail the technical orchestrability assessment come off the list before sequences are built. Companies that clear all three dimensions get prioritized outreach with seller-ready context already in HubSpot.")
    bullet_bold(doc, "The Inspector Case Board as the standard pre-call research tool for all sellers and SEs. ", "The \u2018Run Inspector\u2019 link on every Company and Deal record makes it a one-click motion. The Case Board gives every seller account intelligence they currently do not have before the first conversation.")
    bullet_bold(doc, "Designer as a standard deliverable in every new customer engagement. ", "Skillable LC and PS should run every new customer through Designer in the first week \u2014 while the technical folks are working to get the technical details sorted. Program design and environment build run in parallel.", after=60)

    h2(doc, "Decisions required from RevOps and Marketing")
    body(doc, "The following require RevOps and Marketing input before the HubSpot integration can be built:")
    add_table(doc,
        ["Decision", "Context"],
        [
            ["Existing custom Company properties",   "Which recommended fields already exist vs. net-new?"],
            ["Deal template review",                 "Where does each data element fit in the existing deal UX?"],
            ["Deduplication rules",                  "What constitutes a match to an existing Deal for expansion opportunities?"],
            ["Dossier multi-product question",        "Three products in one Inspector run — three Deals or one?"],
            ["Buying Group Summary structure",        "Current state across the customer base?"],
            ["Ownership and notification rules",     "How should Intelligence-generated Deals trigger notifications for AEs and CSMs?"],
            ["ZoomInfo CSV column mapping",          "Minimum: Company Name + Domain. High value: Industry, LinkedIn URL, Employee Count, Technologies Used."],
            ["Score threshold for auto-create/update", "At what fit score does a Company record get automatically enriched vs. queued for review?"],
        ],
        col1_dxa=3800
    )


# ── Main ──────────────────────────────────────────────────────────────────────

def main():
    doc, section = setup_document()
    add_header(section)
    add_footer(section)

    add_title(doc, "Skillable Intelligence Platform", "Executive Leadership Team Brief")

    write_why(doc)
    write_what(doc)
    write_how(doc)
    write_recommendations(doc)

    doc.save(OUTPUT_PATH)
    print(f"Saved: {OUTPUT_PATH}")


if __name__ == "__main__":
    main()
