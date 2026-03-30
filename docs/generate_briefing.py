#!/usr/bin/env python3
"""
Generate Skillable Intelligence Executive Briefing Word document.
Compact 2-3 page format with logo header, page X of Y footer.
"""

import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Twips, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from lxml import etree

DOCS_DIR = os.path.dirname(os.path.abspath(__file__))
OUTPUT_PATH = os.path.join(DOCS_DIR, "Skillable-Intelligence-Executive-Briefing.docx")
LOGO_PATH = r"C:\Users\Frank.Gartland\OneDrive - Skillable\Sales Enablement\Keep\z-Skillable Logos\Skillable Logo\Default@4x.png"

# Colors
DARK_GREEN = RGBColor(0x1A, 0x4D, 0x35)
DARK_GREEN_HEX = "1A4D35"
GRAY = RGBColor(0x60, 0x60, 0x60)
MEDIUM_GRAY = RGBColor(0xA0, 0xA0, 0xA0)
PAGE_NUM_GRAY = RGBColor(0x88, 0x88, 0x88)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
DARK_TEXT = RGBColor(0x1A, 0x1A, 0x1A)
ROW_ALT = "F0F5F2"
FONT_NAME = "Calibri"


# ── XML helpers ────────────────────────────────────────────────────────────────

def set_cell_background(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def set_cell_borders(cell, color_hex="CCCCCC", size=4):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in ["top", "left", "bottom", "right"]:
        border = OxmlElement(f"w:{side}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), str(size))
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), color_hex)
        tcBorders.append(border)
    tcPr.append(tcBorders)


def set_cell_margins(cell, val=100):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement("w:tcMar")
    for side in ["top", "left", "bottom", "right"]:
        m = OxmlElement(f"w:{side}")
        m.set(qn("w:w"), str(val))
        m.set(qn("w:type"), "dxa")
        tcMar.append(m)
    tcPr.append(tcMar)


def add_paragraph_border_bottom(paragraph, color_hex=DARK_GREEN_HEX, size=8):
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color_hex)
    pBdr.append(bottom)
    pPr.append(pBdr)


def set_paragraph_spacing(paragraph, before=0, after=0, line=None, line_rule=None):
    pPr = paragraph._p.get_or_add_pPr()
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:before"), str(before))
    spacing.set(qn("w:after"), str(after))
    if line is not None:
        spacing.set(qn("w:line"), str(line))
    if line_rule is not None:
        spacing.set(qn("w:lineRule"), line_rule)
    pPr.append(spacing)


def make_run(paragraph, text, size_pt=10, bold=False, italic=False, color=None):
    run = paragraph.add_run(text)
    run.font.name = FONT_NAME
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color
    return run


# ── Content helpers ────────────────────────────────────────────────────────────

def add_h1(doc, text):
    p = doc.add_paragraph()
    r = make_run(p, text, size_pt=12, bold=True, color=DARK_GREEN)
    set_paragraph_spacing(p, before=180, after=60)
    add_paragraph_border_bottom(p, DARK_GREEN_HEX, size=6)
    return p


def add_body(doc, text, space_before=0, space_after=80):
    p = doc.add_paragraph()
    make_run(p, text, size_pt=10, color=DARK_TEXT)
    set_paragraph_spacing(p, before=space_before, after=space_after)
    return p


def add_bullet(doc, text, bold_prefix=None, space_after=60):
    """Add a bullet. If bold_prefix given, render it bold then rest normal."""
    p = doc.add_paragraph(style="List Bullet")
    if bold_prefix:
        make_run(p, bold_prefix, size_pt=10, bold=True, color=DARK_TEXT)
        make_run(p, text, size_pt=10, color=DARK_TEXT)
    else:
        make_run(p, text, size_pt=10, color=DARK_TEXT)
    set_paragraph_spacing(p, before=0, after=space_after)
    return p


def add_tiny_space(doc, pt=4):
    p = doc.add_paragraph()
    p.add_run("").font.size = Pt(pt)
    set_paragraph_spacing(p, before=0, after=0)
    return p


# ── Header / Footer XML ────────────────────────────────────────────────────────

def build_header(doc, logo_path):
    """Return a Header object: Skillable logo left-aligned, dark green rule below."""
    from docx.oxml.ns import nsmap
    from docx.opc.part import Part

    section = doc.sections[0]
    header = section.header
    header.is_linked_to_previous = False

    # Clear default paragraph
    for p in header.paragraphs:
        p._element.getparent().remove(p._element)

    # Logo paragraph
    logo_p = header.add_paragraph()
    logo_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_paragraph_spacing(logo_p, before=0, after=40)

    run = logo_p.add_run()
    with open(logo_path, "rb") as f:
        run.add_picture(f, width=Inches(1.1))

    # Rule paragraph (bottom border only)
    rule_p = header.add_paragraph()
    rule_p.add_run("").font.size = Pt(2)
    set_paragraph_spacing(rule_p, before=0, after=0)
    add_paragraph_border_bottom(rule_p, DARK_GREEN_HEX, size=6)

    return header


def build_footer(doc):
    """Footer: 'Page X of Y' right-aligned in 9pt gray."""
    section = doc.sections[0]
    footer = section.footer
    footer.is_linked_to_previous = False

    for p in footer.paragraphs:
        p._element.getparent().remove(p._element)

    fp = footer.add_paragraph()
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_paragraph_spacing(fp, before=0, after=0)

    def gray_run_xml(text):
        r = OxmlElement("w:r")
        rpr = OxmlElement("w:rPr")
        color = OxmlElement("w:color")
        color.set(qn("w:val"), "888888")
        sz = OxmlElement("w:sz")
        sz.set(qn("w:val"), "16")  # 8pt
        szCs = OxmlElement("w:szCs")
        szCs.set(qn("w:val"), "16")
        font = OxmlElement("w:rFonts")
        font.set(qn("w:ascii"), "Calibri")
        font.set(qn("w:hAnsi"), "Calibri")
        rpr.append(font)
        rpr.append(sz)
        rpr.append(szCs)
        rpr.append(color)
        r.append(rpr)
        t = OxmlElement("w:t")
        t.text = text
        if text.startswith(" ") or text.endswith(" "):
            t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
        r.append(t)
        return r

    def page_num_fld(fld_type):
        """Return an w:fldChar + instrText sequence for PAGE or NUMPAGES."""
        r1 = OxmlElement("w:r")
        rpr1 = OxmlElement("w:rPr")
        color1 = OxmlElement("w:color")
        color1.set(qn("w:val"), "888888")
        sz1 = OxmlElement("w:sz")
        sz1.set(qn("w:val"), "18")
        szCs1 = OxmlElement("w:szCs")
        szCs1.set(qn("w:val"), "18")
        font1 = OxmlElement("w:rFonts")
        font1.set(qn("w:ascii"), "Arial")
        font1.set(qn("w:hAnsi"), "Arial")
        rpr1.append(font1)
        rpr1.append(sz1)
        rpr1.append(szCs1)
        rpr1.append(color1)
        r1.append(rpr1)
        fldChar_begin = OxmlElement("w:fldChar")
        fldChar_begin.set(qn("w:fldCharType"), "begin")
        r1.append(fldChar_begin)

        r2 = OxmlElement("w:r")
        rpr2 = OxmlElement("w:rPr")
        color2 = OxmlElement("w:color")
        color2.set(qn("w:val"), "888888")
        sz2 = OxmlElement("w:sz")
        sz2.set(qn("w:val"), "18")
        szCs2 = OxmlElement("w:szCs")
        szCs2.set(qn("w:val"), "18")
        font2 = OxmlElement("w:rFonts")
        font2.set(qn("w:ascii"), "Arial")
        font2.set(qn("w:hAnsi"), "Arial")
        rpr2.append(font2)
        rpr2.append(sz2)
        rpr2.append(szCs2)
        rpr2.append(color2)
        r2.append(rpr2)
        instr = OxmlElement("w:instrText")
        instr.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
        instr.text = f" {fld_type} "
        r2.append(instr)

        r3 = OxmlElement("w:r")
        rpr3 = OxmlElement("w:rPr")
        color3 = OxmlElement("w:color")
        color3.set(qn("w:val"), "888888")
        sz3 = OxmlElement("w:sz")
        sz3.set(qn("w:val"), "18")
        szCs3 = OxmlElement("w:szCs")
        szCs3.set(qn("w:val"), "18")
        font3 = OxmlElement("w:rFonts")
        font3.set(qn("w:ascii"), "Arial")
        font3.set(qn("w:hAnsi"), "Arial")
        rpr3.append(font3)
        rpr3.append(sz3)
        rpr3.append(szCs3)
        rpr3.append(color3)
        r3.append(rpr3)
        fldChar_end = OxmlElement("w:fldChar")
        fldChar_end.set(qn("w:fldCharType"), "end")
        r3.append(fldChar_end)

        return r1, r2, r3

    p_elem = fp._p
    p_elem.append(gray_run_xml("Page "))
    for el in page_num_fld("PAGE"):
        p_elem.append(el)
    p_elem.append(gray_run_xml(" of "))
    for el in page_num_fld("NUMPAGES"):
        p_elem.append(el)

    return footer


# ── Tables ─────────────────────────────────────────────────────────────────────

def build_flow_and_deps_table(doc):
    """
    Compact combined section: a small left-side flow diagram + deps table on right,
    implemented as a two-column outer table.
    """
    # Outer table: flow (left) | dep table (right)
    # Content width = 9360 DXA (6.5"). Left col = 2700, right = 6660
    outer = doc.add_table(rows=1, cols=2)
    outer.alignment = WD_TABLE_ALIGNMENT.LEFT
    # Remove all outer borders
    tbl = outer._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    tblBorders = OxmlElement("w:tblBorders")
    for side in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"), "none")
        b.set(qn("w:sz"), "0")
        b.set(qn("w:space"), "0")
        b.set(qn("w:color"), "auto")
        tblBorders.append(b)
    tblPr.append(tblBorders)

    # Full content width: 8.5" page - 0.85" margins × 2 = 6.8" = 9792 DXA
    left_cell = outer.rows[0].cells[0]
    right_cell = outer.rows[0].cells[1]
    left_cell.width = Twips(2160)
    right_cell.width = Twips(7632)

    for cell in [left_cell, right_cell]:
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcBorders = OxmlElement("w:tcBorders")
        for side in ["top", "left", "bottom", "right"]:
            b = OxmlElement(f"w:{side}")
            b.set(qn("w:val"), "none")
            b.set(qn("w:sz"), "0")
            b.set(qn("w:space"), "0")
            b.set(qn("w:color"), "auto")
            tcBorders.append(b)
        tcPr.append(tcBorders)

    # ── Left cell: flow diagram ──────────────────────────────────────────

    def lp(text, bold=False, color=None, align=WD_ALIGN_PARAGRAPH.LEFT, size=9):
        p = left_cell.add_paragraph()
        p.alignment = align
        r = p.add_run(text)
        r.font.name = FONT_NAME
        r.font.size = Pt(size)
        r.font.bold = bold
        if color:
            r.font.color.rgb = color
        else:
            r.font.color.rgb = DARK_TEXT
        set_paragraph_spacing(p, before=0, after=30)
        return p

    # Remove default empty paragraph
    for p in list(left_cell.paragraphs):
        p._element.getparent().remove(p._element)

    def flow_box(cell_obj, text, bg_hex, text_color=WHITE):
        p = cell_obj.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r = p.add_run(text)
        r.font.name = FONT_NAME
        r.font.size = Pt(8)
        r.font.bold = True
        r.font.color.rgb = text_color
        set_paragraph_spacing(p, before=20, after=20)
        # Add shading to paragraph via paragraph shade
        pPr = p._p.get_or_add_pPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), bg_hex)
        pPr.append(shd)
        # Indent
        ind = OxmlElement("w:ind")
        ind.set(qn("w:left"), "120")
        ind.set(qn("w:right"), "120")
        pPr.append(ind)
        return p

    def flow_arrow(cell_obj, char="\u2193"):
        p = cell_obj.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r = p.add_run(char)
        r.font.name = FONT_NAME
        r.font.size = Pt(8)
        r.font.bold = True
        r.font.color.rgb = DARK_GREEN
        set_paragraph_spacing(p, before=0, after=0)
        ind = OxmlElement("w:ind")
        ind.set(qn("w:left"), "160")
        p._p.get_or_add_pPr().append(ind)
        return p

    # Label
    p_label = left_cell.add_paragraph()
    r_label = p_label.add_run("Pipeline Flow")
    r_label.font.name = FONT_NAME
    r_label.font.size = Pt(8)
    r_label.font.bold = True
    r_label.font.color.rgb = GRAY
    set_paragraph_spacing(p_label, before=0, after=60)

    flow_box(left_cell, "ZoomInfo Export", "D0E8DA", DARK_TEXT)
    flow_arrow(left_cell)
    flow_box(left_cell, "PROSPECTOR", DARK_GREEN_HEX)
    flow_arrow(left_cell)
    flow_box(left_cell, "INSPECTOR  \u2192  HubSpot", DARK_GREEN_HEX)
    flow_arrow(left_cell)
    flow_box(left_cell, "DESIGNER", DARK_GREEN_HEX)
    flow_arrow(left_cell)
    flow_box(left_cell, "Skillable Studio", "D0E8DA", DARK_TEXT)

    # ── Right cell: dependency table ─────────────────────────────────────

    # Remove default empty paragraph
    for p in list(right_cell.paragraphs):
        p._element.getparent().remove(p._element)

    # Label
    p_dep_label = right_cell.add_paragraph()
    r_dep_label = p_dep_label.add_run("External Dependencies")
    r_dep_label.font.name = FONT_NAME
    r_dep_label.font.size = Pt(8)
    r_dep_label.font.bold = True
    r_dep_label.font.color.rgb = GRAY
    set_paragraph_spacing(p_dep_label, before=0, after=60)

    dep_headers = ["Dependency", "Purpose", "Owner"]
    dep_rows = [
        ["Anthropic Claude API", "AI research & scoring \u2014 all three tools", "Platform (Key Vault)"],
        ["Serper.dev (Google Search)", "Web research for company/product data", "Platform (Key Vault)"],
        ["ZoomInfo", "Source of company lists for Prospector", "Revenue Operations"],
        ["HubSpot", "Destination for qualified leads & opportunity data", "Revenue Operations"],
    ]
    # Col widths inside right cell (6840 DXA total, minus small left pad)
    # Dep table fills right cell (7632 DXA): Dependency | Purpose | Owner
    dep_col_w = [2200, 3832, 1600]

    dep_tbl = right_cell.add_table(rows=1 + len(dep_rows), cols=3)
    dep_tbl.alignment = WD_TABLE_ALIGNMENT.LEFT

    def dep_cell(row, col, text, header=False, alt=False):
        cell = row.cells[col]
        cell.width = Twips(dep_col_w[col])
        p = cell.paragraphs[0]
        r = p.add_run(text)
        r.font.name = FONT_NAME
        r.font.size = Pt(8)
        r.font.bold = header
        if header:
            r.font.color.rgb = WHITE
            set_cell_background(cell, DARK_GREEN_HEX)
            set_cell_borders(cell, DARK_GREEN_HEX, size=4)
        else:
            r.font.color.rgb = DARK_TEXT
            set_cell_background(cell, ROW_ALT if alt else "FFFFFF")
            set_cell_borders(cell, "CCCCCC", size=4)
        set_cell_margins(cell, 80)

    for i, h in enumerate(dep_headers):
        dep_cell(dep_tbl.rows[0], i, h, header=True)
    for ri, row_data in enumerate(dep_rows):
        row = dep_tbl.rows[ri + 1]
        for ci, txt in enumerate(row_data):
            dep_cell(row, ci, txt, alt=(ri % 2 == 0))


# ── Two-column Open Questions ──────────────────────────────────────────────────

def build_open_questions_2col(doc):
    """
    Open questions rendered as a borderless 2-column table:
    SecOps on left, RevOps on right.
    """
    outer = doc.add_table(rows=1, cols=2)
    outer.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Remove borders from the outer table
    tbl = outer._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    tblBorders = OxmlElement("w:tblBorders")
    for side in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"), "none")
        b.set(qn("w:sz"), "0")
        b.set(qn("w:space"), "0")
        b.set(qn("w:color"), "auto")
        tblBorders.append(b)
    tblPr.append(tblBorders)

    left_cell = outer.rows[0].cells[0]
    right_cell = outer.rows[0].cells[1]
    left_cell.width = Twips(4500)
    right_cell.width = Twips(4860)

    for cell in [left_cell, right_cell]:
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcBorders = OxmlElement("w:tcBorders")
        for side in ["top", "left", "bottom", "right"]:
            b = OxmlElement(f"w:{side}")
            b.set(qn("w:val"), "none")
            b.set(qn("w:sz"), "0")
            b.set(qn("w:space"), "0")
            b.set(qn("w:color"), "auto")
            tcBorders.append(b)
        tcPr.append(tcBorders)
        set_cell_margins(cell, 0)

    def add_col_heading(cell, text):
        for p in list(cell.paragraphs):
            p._element.getparent().remove(p._element)
        p = cell.add_paragraph()
        r = p.add_run(text)
        r.font.name = FONT_NAME
        r.font.size = Pt(10)
        r.font.bold = True
        r.font.color.rgb = DARK_GREEN
        set_paragraph_spacing(p, before=0, after=60)

    def add_col_bullet(cell, text):
        p = cell.add_paragraph(style="List Bullet")
        r = p.add_run(text)
        r.font.name = FONT_NAME
        r.font.size = Pt(9)
        r.font.color.rgb = DARK_TEXT
        set_paragraph_spacing(p, before=0, after=50)

    def add_col_bullet_prefixed(cell, bold_prefix, text):
        p = cell.add_paragraph(style="List Bullet")
        r_bold = p.add_run(bold_prefix)
        r_bold.font.name = FONT_NAME
        r_bold.font.size = Pt(9)
        r_bold.font.bold = True
        r_bold.font.color.rgb = DARK_TEXT
        r_rest = p.add_run(text)
        r_rest.font.name = FONT_NAME
        r_rest.font.size = Pt(9)
        r_rest.font.color.rgb = DARK_TEXT
        set_paragraph_spacing(p, before=0, after=50)

    add_col_heading(left_cell, "SecOps")
    add_col_bullet(left_cell, "What\u2019s the process for moving a repo to the corporate GitHub org? Who approves access and sets permissions?")
    add_col_bullet(left_cell, "Which Azure subscription and resource group should this live in, and what\u2019s the approval process for deploying a new internal tool?")
    add_col_bullet(left_cell, "Standard approach for Entra ID app registration for internal tools?")
    add_col_bullet(left_cell, "Data classification for JSON analysis files (public company info + LinkedIn URLs)?")
    add_col_bullet(left_cell, "DLP or egress controls that could affect outbound calls to Anthropic (US) or Serper APIs?")
    add_col_bullet(left_cell, "Do you require audit logging for internal tools that make external API calls? If so, what\u2019s the standard \u2014 we\u2019d log user, company analyzed, timestamp, and APIs called.")

    add_col_heading(right_cell, "RevOps")
    add_col_bullet_prefixed(right_cell, "Recommended:", " Add a \u201cProducts\u201d card to the HubSpot Company record \u2014 each product listed as a link to its Inspector score page, labeled Highly Labable / Likely Labable / Not Labable.")
    add_col_bullet_prefixed(right_cell, "Recommended:", " Add a \u201cLab Maturity Signals \u2192\u201d link on the Company record pointing to the full lab maturity breakdown page in Inspector.")
    add_col_bullet_prefixed(right_cell, "Recommended:", " Prospector results should update existing HubSpot Contacts or add new ones \u2014 not a separate object.")
    add_col_bullet_prefixed(right_cell, "Recommended:", " RevOps owns the HubSpot integration build.")
    add_col_bullet_prefixed(right_cell, "Recommended:", " Identify the existing HubSpot source attribution field (\u201cLead Source\u201d or equivalent) and confirm the stamp value (e.g., \u201cSkillable Intelligence\u201d) \u2014 enables pipeline influence reporting.")
    add_col_bullet(right_cell, "Standard ZoomInfo export format? Which filters to apply before feeding Prospector?")
    add_col_bullet(right_cell, "What other tools are in the RevOps stack (Salesforce, Outreach, Apollo, Sales Navigator, etc.) that should be considered for integration or data handoff?")


# ── Main ───────────────────────────────────────────────────────────────────────

def main():
    if not os.path.exists(LOGO_PATH):
        raise FileNotFoundError(f"Logo not found at: {LOGO_PATH}")

    doc = Document()

    # Page setup: US Letter, tight margins (0.75" sides, 1" top/bottom)
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)
    section.top_margin = Inches(0.9)
    section.bottom_margin = Inches(0.9)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)

    # Default font
    doc.styles["Normal"].font.name = FONT_NAME
    doc.styles["Normal"].font.size = Pt(10)

    # Header and footer
    build_header(doc, LOGO_PATH)
    build_footer(doc)

    # ── TITLE BLOCK ───────────────────────────────────────────────────────

    title_p = doc.add_paragraph()
    t_run = title_p.add_run("Skillable Intelligence Platform")
    t_run.font.name = FONT_NAME
    t_run.font.size = Pt(20)
    t_run.font.bold = True
    t_run.font.color.rgb = DARK_GREEN
    set_paragraph_spacing(title_p, before=0, after=40)

    sub_p = doc.add_paragraph()
    s_run = sub_p.add_run("Executive Briefing \u2014 Security Operations & Revenue Operations  \u00b7  March 2026")
    s_run.font.name = FONT_NAME
    s_run.font.size = Pt(10)
    s_run.font.color.rgb = GRAY
    set_paragraph_spacing(sub_p, before=0, after=60)

    rule_p = doc.add_paragraph()
    rule_p.add_run("").font.size = Pt(2)
    set_paragraph_spacing(rule_p, before=0, after=140)
    add_paragraph_border_bottom(rule_p, DARK_GREEN_HEX, size=10)

    # ── SECTION 1: Why We Built This ─────────────────────────────────────

    add_h1(doc, "1.  Why We Built This \u2014 Problems Being Solved")

    add_body(doc,
        "The revenue and partnerships teams face three concrete friction points today:",
        space_after=60)

    add_bullet(doc,
        " No scalable way to build target lists. Identifying companies with the right training ecosystems, technical products, and partner programs for a Skillable conversation currently requires manual research at every stage.",
        bold_prefix="Targeting & Lead Qualification:")
    add_bullet(doc,
        " No fast way to qualify a company once identified. SEs and AMs spend hours researching whether a prospect\u2019s products are technically buildable as Skillable labs \u2014 time that compounds across every territory.",
        bold_prefix="Accelerate Discovery:")
    add_bullet(doc,
        " No starting point for lab program design. When an SE wants to propose a lab program, they start from a blank page \u2014 no structure, no connection to the research already done on the company.",
        bold_prefix="Facilitate Lab Design & Decrease TTV:")

    closing = add_body(doc, "Skillable Intelligence addresses all three.", space_before=40, space_after=100)
    closing.runs[0].font.bold = True

    # ── SECTION 2: What It Is ─────────────────────────────────────────────

    add_h1(doc, "2.  What It Is \u2014 Platform Elements")

    add_body(doc,
        "Three tools forming a pipeline from target identification through lab program design:",
        space_after=60)

    add_bullet(doc,
        " Runs a lightweight Inspector analysis on a full ZoomInfo export simultaneously \u2014 returns a ranked, scored table with labability/lab maturity/composite scores, a path label (Labable / Simulations / Do Not Pursue / Academic / Not a Fit), and top two contacts per company with LinkedIn and Excel export. Used by RevOps and field sales for account-based targeting.",
        bold_prefix="Prospector \u2014 Batch Territory Scoring.")
    add_bullet(doc,
        " Scores a company\u2019s products for labability (0\u2013100) across four dimensions; produces a lab maturity score, opportunity composite, recommended Skillable path (VM / Cloud Slice / Simulation), estimated annual ACV, and named contacts for training/enablement outreach. Used by SEs and AMs for pre-call research.",
        bold_prefix="Inspector \u2014 Single-Company Deep Analysis.")
    add_bullet(doc,
        " A four-phase wizard that designs a complete lab program for a specific company \u2014 from audience and objectives through program architecture, draft lab instructions, and Skillable Studio export. Used by SEs building a proof of concept or proposal after Inspector identifies a strong opportunity.",
        bold_prefix="Designer \u2014 Lab Program Scaffolding.")

    # ── SECTION 3: How They Work Together ────────────────────────────────

    add_h1(doc, "3.  How They Work Together")
    add_tiny_space(doc, 4)
    build_flow_and_deps_table(doc)
    add_tiny_space(doc, 6)

    # ── SECTION 4: Security ───────────────────────────────────────────────

    add_h1(doc, "4.  Security \u2014 What SecOps Needs to Know")

    add_bullet(doc,
        " Two API keys in a server-side .env file \u2014 never in browser storage or source code. (1) Anthropic Claude API \u2014 AI research and scoring engine; all three tools depend on it. (2) Serper.dev \u2014 Google Search API used for web research on every company analysis; requires a paid subscription. Both keys move to Azure Key Vault in production.",
        bold_prefix="Credential Handling:")
    add_bullet(doc,
        " Company names, URLs, and product descriptions (public, AI-synthesized). Contact names, titles, and LinkedIn URLs sourced from public web search \u2014 equivalent to a manual Google search. No customer data, financial data, or internal Skillable systems data flows through the platform.",
        bold_prefix="Data Touched:")
    add_bullet(doc,
        " No user authentication \u2014 currently runs on a developer\u2019s local machine. No database; analysis results stored as flat JSON files on the local filesystem. Results are aggressively cached: re-analyzing a company already in the cache skips all Anthropic and Serper API calls, reducing cost and latency significantly.",
        bold_prefix="Current State \u2014 Storage & Caching:")
    add_bullet(doc,
        " API key exposure (low risk today, needs Azure Key Vault in production). Unauthenticated access (needs Entra ID/SSO before any shared deployment). Data residency: AI calls to Anthropic (US), web search via Serper/Google, no internal data egress. Repo currently on personal GitHub account \u2014 needs to move to corporate Skillable org. No audit trail of who ran which analyses.",
        bold_prefix="Risks:")

    # Directory tree
    add_tiny_space(doc, 6)
    tree_label = doc.add_paragraph()
    tree_label.add_run("Current repo structure (github.com/fgartland4/skillable-intelligence):").font.size = Pt(8)
    tree_label.runs[0].font.name = FONT_NAME
    tree_label.runs[0].font.color.rgb = GRAY
    set_paragraph_spacing(tree_label, before=40, after=40)

    TREE_LINES = [
        "skillable-intelligence/          \u2190 Personal GitHub (fgartland4) \u2014 needs to move to Skillable org",
        "\u251c\u2500\u2500 backend/                     Flask app: routes, scoring engine, AI calls, storage",
        "\u251c\u2500\u2500 tools/",
        "\u2502   \u251c\u2500\u2500 inspector/templates/     Inspector UI",
        "\u2502   \u251c\u2500\u2500 prospector/templates/    Prospector UI",
        "\u2502   \u2514\u2500\u2500 designer/                Designer UI + client JS",
        "\u251c\u2500\u2500 static/                      Shared CSS, JS, images",
        "\u251c\u2500\u2500 docs/                        Executive briefing + generator",
        "\u251c\u2500\u2500 .gitignore",
        "\u2514\u2500\u2500 render.yaml                  Deployment config",
    ]
    for line in TREE_LINES:
        p = doc.add_paragraph()
        r = p.add_run(line)
        r.font.name = "Consolas"
        r.font.size = Pt(8)
        r.font.color.rgb = DARK_TEXT
        set_paragraph_spacing(p, before=0, after=0)
        pPr = p._p.get_or_add_pPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), "F2F5F3")
        pPr.append(shd)
        ind = OxmlElement("w:ind")
        ind.set(qn("w:left"), "200")
        ind.set(qn("w:right"), "200")
        pPr.append(ind)
    add_tiny_space(doc, 6)

    # ── SECTION 5: How We're Going to Do This ────────────────────────────

    add_h1(doc, "5.  How We\u2019re Going to Do This")

    add_bullet(doc,
        " Move repo to corporate GitHub. Add Azure Key Vault (replace .env). Add Entra ID/SSO. Define App Service environment (region, SKU, internal-only). Deploy and confirm HTTPS + access controls.",
        bold_prefix="Phase 1 \u2014 Secure the Foundation (SecOps + Engineering):")
    add_bullet(doc,
        " Define ZoomInfo export \u2192 Prospector import spec. Define what Prospector and Inspector push to HubSpot: Products card (product name, labability tier, link to product detail page), Lab Maturity score, composite score, and top contacts. Build HubSpot integration (Excel import as MVP, API as next step). Run pilot: one territory through Prospector \u2192 HubSpot.",
        bold_prefix="Phase 2 \u2014 Connect Revenue Workflows (RevOps + Engineering):")
    add_bullet(doc,
        " Instrument usage (companies analyzed, score distributions, pipeline conversion). Complete Designer Phase 2 and 3 prompts and Studio handoff format. Upgrade search API once usage exceeds ~50 Inspector analyses/month. Expand customer benchmarks as new logos close.",
        bold_prefix="Phase 3 \u2014 Expand and Iterate:")

    # ── SECTION 6: Open Questions ─────────────────────────────────────────

    add_h1(doc, "6.  Open Questions")
    add_tiny_space(doc, 4)
    build_open_questions_2col(doc)

    # ── Save ──────────────────────────────────────────────────────────────
    doc.save(OUTPUT_PATH)
    size_kb = os.path.getsize(OUTPUT_PATH) / 1024
    print(f"Saved: {OUTPUT_PATH}")
    print(f"File size: {size_kb:.1f} KB")


if __name__ == "__main__":
    main()
